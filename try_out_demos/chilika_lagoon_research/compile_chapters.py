import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

base_path = "C:/Users/Dancy Naik/Documents/VS_Code_Test/wow_ai/try_out_demos/chilika_lagoon_research/"
chapters_path = os.path.join(base_path, "chapters")
output_file = os.path.join(base_path, "Chilika_Water_Quality_Review.docx")

# Files to read
chapter_files = [
    f"chapter_{i}_{title}.md" for i, title in enumerate([
        "introduction",
        "background_literature_review",
        "data_collection_and_methods",
        "analysis_results",
        "discussion",
        "management_implications",
        "conclusion_references"
    ], start=1)
]

# APA 7th edition basics for references section (e.g., hanging indent, font style/size set)
# We will keep overall style with font Times New Roman, size 12, double-spaced.


def add_page_number(doc):
    # Add page number to footer of the first section
    section = doc.sections[0]
    footer = section.footer
    if not footer.paragraphs:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    r = p.add_run()
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)


def add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update Table of Contents"

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)


def set_paragraph_style(paragraph, font_name='Times New Roman', font_size=12, bold=False, italic=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, space_after=12, space_before=12, double_space=True):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic
    paragraph.alignment = alignment
    paragraph.space_after = Pt(space_after)
    paragraph.space_before = Pt(space_before)
    if double_space:
        paragraph.line_spacing = 2.0


def markdown_to_text(markdown_content):
    # Basic markdown to plain text, supports # headings (single level), bold (**text**), italic (*text*)
    # Does not support complex markdown, simplified for our use case.
    lines = markdown_content.splitlines()
    parsed_lines = []
    for line in lines:
        if line.startswith('# '):
            parsed_lines.append(line[2:].strip())
        elif line.startswith('## '):
            parsed_lines.append(line[3:].strip())
        else:
            # Replace bold and italic markdown
            line = line.replace('**', '')
            line = line.replace('*', '')
            parsed_lines.append(line)
    return "\n".join(parsed_lines)


def main():
    doc = Document()

    # Title page
    title = "Chilika Water Quality Review"
    author = "Compiled Research"

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # blank

    author_paragraph = doc.add_paragraph()
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_a = author_paragraph.add_run(author)
    run_a.font.size = Pt(14)
    run_a.font.name = 'Times New Roman'

    doc.add_page_break()

    # Table of contents placeholder
    toc_paragraph = doc.add_paragraph('Table of Contents')
    toc_paragraph.style = 'Heading 1'
    add_table_of_contents(doc)

    doc.add_page_break()

    # Read and add chapters with headings
    for i, chapter_file in enumerate(chapter_files, start=1):
        chapter_path = os.path.join(chapters_path, chapter_file)
        if not os.path.exists(chapter_path):
            print(f"Warning: Missing file {chapter_path}")
            continue
        with open(chapter_path, 'r', encoding='utf-8') as f:
            content = f.read()

        heading_text = f"Chapter {i}: " + chapter_file.split('_', 2)[2].replace('.md', '').replace('_', ' ').title()
        chapter_heading = doc.add_paragraph(heading_text, style='Heading 1')

        # Add content paragraphs
        chapter_text = markdown_to_text(content)
        for para_text in chapter_text.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph(para_text)
            set_paragraph_style(p)

        doc.add_page_break()

    # Add page numbering
    add_page_number(doc)

    # Save Word document
    doc.save(output_file)
    print(f"Document compiled and saved to {output_file}")


if __name__ == '__main__':
    main()
