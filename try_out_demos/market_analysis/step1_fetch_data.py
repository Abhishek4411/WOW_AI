import subprocess
import sys
import os

# Function to install packages
def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Install required libraries
required_packages = ['yfinance', 'pandas', 'matplotlib', 'python-docx']
for package in required_packages:
    try:
        __import__(package)
    except ImportError:
        install(package)

import yfinance as yf
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Define file paths
output_dir = os.path.dirname(os.path.abspath(__file__))
chart_path = os.path.join(output_dir, 'moving_averages.png')
report_path = os.path.join(output_dir, 'Stock_Performance_Report.docx')

# Fetch daily price data for AAPL and MSFT for the year 2023
start_date = '2023-01-01'
end_date = '2023-12-31'
tickers = ['AAPL', 'MSFT']

data = yf.download(tickers, start=start_date, end=end_date, group_by='ticker', progress=False)

# Calculate daily volatility (daily returns std dev) and 30-day moving average (closing price)
for ticker in tickers:
    if 'Adj Close' in data[ticker].columns:
        close_col = 'Adj Close'
    else:
        close_col = 'Close'
    data[ticker, 'Daily_Return'] = data[ticker][close_col].pct_change()
    data[ticker, 'Volatility'] = data[ticker, 'Daily_Return'].rolling(window=1).std()
    data[ticker, '30d_MA'] = data[ticker][close_col].rolling(window=30).mean()

# Plot 30-day moving averages for both stocks
plt.figure(figsize=(12, 6))
plt.plot(data['AAPL']['30d_MA'], label='AAPL 30-Day MA')
plt.plot(data['MSFT']['30d_MA'], label='MSFT 30-Day MA')
plt.title('30-Day Moving Averages of AAPL and MSFT in 2023')
plt.xlabel('Date')
plt.ylabel('Price (USD)')
plt.legend()
plt.grid(True)
plt.tight_layout()
plt.savefig(chart_path)
plt.close()

# Prepare the Word document report

doc = Document()
doc.add_heading('Stock Performance Report: AAPL vs MSFT in 2023', 0)

doc.add_heading('Methodology', level=1)
doc.add_paragraph(
    "This report analyzes the daily stock price performance of Apple Inc. (AAPL) and Microsoft Corporation (MSFT) for the calendar year 2023. "
    "The data was fetched using the yfinance library, obtaining daily adjusted closing prices. "
    "Daily volatility was calculated by observing the standard deviation of daily returns, and a 30-day moving average was computed to smooth out short-term fluctuations in stock prices."
)

# Insert the chart into the report
doc.add_heading('30-Day Moving Average Comparison', level=1)
doc.add_paragraph("Figure 1 below shows the 30-day moving averages of AAPL and MSFT throughout 2023.")
doc.add_picture(chart_path, width=Inches(6))

# Calculate overall volatility for the year to conclude stability
aapl_volatility = data['AAPL']['Daily_Return'].std()
msft_volatility = data['MSFT']['Daily_Return'].std()

conclusion = "Based on the calculated daily return volatilities, "
if aapl_volatility < msft_volatility:
    conclusion += "Apple (AAPL) demonstrated slightly more stable stock performance than Microsoft (MSFT) in 2023."
else:
    conclusion += "Microsoft (MSFT) demonstrated slightly more stable stock performance than Apple (AAPL) in 2023."

# Add conclusion section to the report
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(conclusion)

doc.save(report_path)

print(f"Report and chart saved to {output_dir}")
